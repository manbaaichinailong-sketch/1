from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUTPUT_DIR = Path("output")
OUTPUT_DIR.mkdir(exist_ok=True)
OUTPUT_FILE = OUTPUT_DIR / "AI干散货航次利润动态预测系统_教师展示简版.docx"

NAVY = "17365D"
BLUE = "1F4E78"
LIGHT_BLUE = "EAF3F8"
LIGHT_YELLOW = "FFF8E1"
GRID = "B7C9D6"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text_color(cell, color, bold=False):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor.from_string(color)
            run.bold = bold


def set_table_borders(table, color=GRID, size="4"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_run_font(run, size=11, bold=False, color=None, name="Microsoft YaHei"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(text)
    if level == 1:
        set_run_font(run, 15, True, BLUE)
        p.paragraph_format.keep_with_next = True
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "4")
        bottom.set(qn("w:color"), "A9C4DB")
        pBdr.append(bottom)
        pPr.append(pBdr)
    else:
        set_run_font(run, 12, True, NAVY)
    return p


def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.45
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.first_line_indent = Cm(0.74)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, 11, True, NAVY)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2, 11)
    else:
        run = p.add_run(text)
        set_run_font(run, 11)
    return p


def add_bullet(doc, text, numbered=False):
    style = "List Number" if numbered else "List Bullet"
    p = doc.add_paragraph(style=style)
    p.paragraph_format.line_spacing = 1.3
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_run_font(run, 10.5)
    return p


def add_info_box(doc, label, text, fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_table_borders(table, color="9AB7CF", size="6")
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r1 = p.add_run(label)
    set_run_font(r1, 11, True, NAVY)
    r2 = p.add_run(text)
    set_run_font(r2, 11)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_two_col_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = True
    headers = table.rows[0].cells
    headers[0].text = "模块"
    headers[1].text = "主要内容"
    for cell in headers:
        set_cell_shading(cell, BLUE)
        set_cell_text_color(cell, "FFFFFF", True)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for left, right in rows:
        cells = table.add_row().cells
        cells[0].text = left
        cells[1].text = right
        cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for p in cells[0].paragraphs:
            for r in p.runs:
                set_run_font(r, 10.5, True, NAVY)
        for p in cells[1].paragraphs:
            for r in p.runs:
                set_run_font(r, 10.5)
    set_table_borders(table)
    for row in table.rows:
        row.cells[0].width = Cm(4.2)
        row.cells[1].width = Cm(11.8)
    return table


def add_stage_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = table.rows[0].cells
    headers[0].text = "阶段"
    headers[1].text = "主要成果"
    for cell in headers:
        set_cell_shading(cell, BLUE)
        set_cell_text_color(cell, "FFFFFF", True)
    for stage, result in rows:
        cells = table.add_row().cells
        cells[0].text = stage
        cells[1].text = result
        set_cell_shading(cells[0], "F3F7FA")
        for p in cells[0].paragraphs:
            for r in p.runs:
                set_run_font(r, 10.5, True, NAVY)
        for p in cells[1].paragraphs:
            for r in p.runs:
                set_run_font(r, 10.5)
    set_table_borders(table)
    return table


def build_document():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2.1)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.2)
    sec.right_margin = Cm(2.2)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(55)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("AI驱动的干散货航次利润动态预测系统")
    set_run_font(r, 22, True, NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(45)
    r = p.add_run("教师沟通简版项目方案")
    set_run_font(r, 13, False, "666666")

    add_info_box(
        doc,
        "项目一句话：",
        "整合AIS船舶动态、燃油价格、港口信息和人工确认的商业数据，通过固定公式动态更新航次预计利润，并利用AI完成参数提取、风险解释和报告生成。",
    )

    add_heading(doc, "一、为什么做这个项目")
    add_body(doc, "干散货航次利润会同时受到运价、货量、航程、船速、油耗、燃油价格、港口费用、等待时间、租金和佣金等因素影响。传统Excel模型通常依赖人工录入，数据更新慢，而且市场参考数据、模型估算数据和最终实际数据容易混在一起。")
    add_body(doc, "本项目希望解决的核心问题是：如何在数据并不完全真实、实时和统一的情况下，持续得到一份来源清楚、可复核、能动态更新的航次利润预测。")

    add_heading(doc, "二、系统由什么组成")
    add_two_col_table(doc, [
        ("数据采集层", "接入VesselFinder、MarineTraffic/Kpler、Spire等专业平台，获取船位、航速、航向、ETA和抵离港事件；同时获取燃油市场价格、天气、港口拥堵及部分公开收费信息。"),
        ("人工数据层", "录入运价、货量、租金、实际油耗、PDA/FDA、拖轮费、引航费、代理费及其他无法公开获取的商业数据。"),
        ("计算引擎", "使用Excel或Python固定公式计算航次天数、燃油成本、港口成本、租金、利润、TCE和保本运价。核心计算不交给AI自由生成。"),
        ("AI分析层", "从自然语言、邮件或业务记录中提取参数，检查缺失项，调用计算引擎，对比不同方案，解释利润变化并生成风险提示和分析报告。"),
    ])
    add_info_box(doc, "核心公式：", "航次利润 ＝ 运费收入 − 佣金 − 燃油成本 − 港口及其他费用 − 租金成本", fill="F7F7F7")

    add_heading(doc, "三、利润结果分三个阶段")
    add_bullet(doc, "Estimated P&L：航次开始前，根据市场油价、预计航程、标准油耗和预计港口费用得到初始估算。", numbered=True)
    add_bullet(doc, "Forecast P&L：航次执行中，根据AIS船位、实际航速、实际等待时间、最新PDA和已确认费用持续更新。", numbered=True)
    add_bullet(doc, "Actual P&L：航次结束后，使用实际燃油发票、FDA、实际租金天数、佣金和滞期费等进行最终核算。", numbered=True)
    add_body(doc, "因此，系统在航次结束前输出的是动态预测利润，而不是声称已经得到最终实际利润。")

    add_heading(doc, "四、怎样解决数据不可靠问题")
    add_body(doc, "每一项数据都保存来源、更新时间、数据类型、可信等级和人工确认状态。缺少来源的数据不进入正式计算，超过规定更新时间的数据自动标记为可能过期。")
    add_info_box(doc, "数据优先级：", "最终账单和发票 ＞ 已确认业务数据/PDA ＞ 专业API数据 ＞ 公开市场数据 ＞ 模型估算数据")
    add_body(doc, "系统不只显示一个利润数字，同时展示利润基准值、可能区间、数据完整度、当前可信度和最大不确定因素。")

    add_heading(doc, "五、华海实习能提供什么")
    add_body(doc, "实习期间主要记录非敏感的业务结构和时间节点，例如ETA、实际到港、靠泊、开工、完工、离港、等泊时间、装卸效率、延误原因、费用类别和代理协调方式。")
    add_body(doc, "不复制企业邮件、合同、客户名称、实际报价和内部账单。用于项目展示时，将船名、客户、日期和金额进行匿名化或比例调整。")

    add_heading(doc, "六、项目实施步骤")
    add_stage_table(doc, [
        ("第一阶段", "完成Excel航次估算模型，计算利润、TCE、保本运价及敏感性分析。"),
        ("第二阶段", "使用Python重建计算引擎，保证公式标准化和可复核。"),
        ("第三阶段", "接入单艘船AIS数据、航线距离和燃油参考价格。"),
        ("第四阶段", "使用Streamlit建立网页界面，展示船舶动态、利润变化和数据可信度。"),
        ("第五阶段", "加入AI参数提取、缺失检查、风险解释和自动报告功能。"),
    ])

    add_heading(doc, "七、预期成果")
    for item in [
        "航次估算Excel模型；",
        "Python航次利润计算程序；",
        "AIS及市场数据接口模块；",
        "航次利润动态预测网页；",
        "AI航次分析助手；",
        "脱敏港口作业案例和研究报告；",
        "完整GitHub项目仓库。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "八、希望向老师请教的问题")
    questions = [
        "学校是否可以提供AIS或专业航运数据库账号？",
        "是否可以申请MarineTraffic、Kpler、Spire、Veson等平台的学术试用？",
        "航次利润模型应采用哪些行业标准和验证方式？",
        "实习数据应如何脱敏，才能兼顾研究价值和企业保密要求？",
        "该项目是否适合作为创新创业、学科竞赛或毕业论文方向？",
    ]
    for q in questions:
        add_bullet(doc, q, numbered=True)

    add_info_box(
        doc,
        "项目定位：",
        "本项目不是让AI“猜利润”，而是建立一个数据来源可追溯、计算公式可核验、预测结果可持续更新的航次经营决策辅助工具。",
        fill=LIGHT_YELLOW,
    )

    # Footer
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("AI干散货航次利润动态预测系统｜教师沟通简版")
        set_run_font(r, 8.5, False, "888888")

    doc.save(OUTPUT_FILE)
    print(OUTPUT_FILE)


if __name__ == "__main__":
    build_document()
